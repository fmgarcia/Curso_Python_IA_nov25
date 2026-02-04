"""
Script para generar documentación en formato DOCX desde archivos Markdown
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_markdown_text(text):
    """Limpia el texto de markdown simple (bold, italic, code)"""
    # Eliminar enlaces [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Eliminar bold/italic
    text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^\*]+)\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    # Eliminar inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

def process_markdown_file(doc, filepath):
    """Procesa un archivo markdown y lo añade al documento"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_lines = []
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.rstrip()
        
        # Bloques de código
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                code_lines = []
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Tablas
        if '|' in line and line.strip() and not line.strip().startswith('<!--'):
            if not in_table:
                in_table = True
                table_data = []
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells and not all(c.replace('-', '').strip() == '' for c in cells):
                table_data.append(cells)
            continue
        else:
            if in_table and table_data:
                # Crear tabla
                if len(table_data) >= 2:
                    headers = table_data[0]
                    rows = [row for row in table_data[1:] if row]
                    
                    if headers and rows:
                        try:
                            table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                            table.style = 'Light Grid Accent 1'
                            
                            # Headers
                            for j, header in enumerate(headers):
                                cell = table.rows[0].cells[j]
                                cell.text = clean_markdown_text(header)
                                if cell.paragraphs:
                                    for run in cell.paragraphs[0].runs:
                                        run.font.bold = True
                            
                            # Rows
                            for row_idx, row_data in enumerate(rows):
                                for col_idx in range(min(len(row_data), len(headers))):
                                    table.rows[row_idx+1].cells[col_idx].text = clean_markdown_text(row_data[col_idx])
                        except:
                            pass
                
                in_table = False
                table_data = []
        
        # Headers
        if line.startswith('#### '):
            doc.add_heading(clean_markdown_text(line[5:]), level=4)
        elif line.startswith('### '):
            doc.add_heading(clean_markdown_text(line[4:]), level=3)
        elif line.startswith('## '):
            doc.add_heading(clean_markdown_text(line[3:]), level=2)
        elif line.startswith('# '):
            doc.add_heading(clean_markdown_text(line[2:]), level=1)
        # Listas
        elif line.startswith('- '):
            p = doc.add_paragraph(clean_markdown_text(line[2:]), style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(clean_markdown_text(text), style='List Number')
        # Línea horizontal
        elif line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 60)
        # Línea vacía
        elif not line.strip():
            if len(doc.paragraphs) > 0:
                doc.add_paragraph()
        # Texto normal
        else:
            if line.strip():
                doc.add_paragraph(clean_markdown_text(line))

def main():
    """Función principal"""
    print("🔨 Generando documentación en formato DOCX...")
    
    try:
        # Crear documento
        doc = Document()
        
        # Configurar el documento
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
        
        # Título principal
        title = doc.add_heading('Sistema de Detección de Tumores', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Centro Médico - Aplicación de Inteligencia Artificial')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if subtitle.runs:
            subtitle.runs[0].font.size = Pt(14)
            subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Orden de archivos
        files_order = [
            'index.md',
            'arquitectura.md',
            'instalacion.md',
            'uso.md',
            'api.md',
            'modelos.md',
            'testing.md'
        ]
        
        # Directorio de documentación
        docs_dir = os.path.dirname(os.path.abspath(__file__))
        
        # Procesar cada archivo
        for filename in files_order:
            filepath = os.path.join(docs_dir, filename)
            if os.path.exists(filepath):
                print(f"  📄 Procesando {filename}...")
                
                # Añadir salto de página antes de cada sección (excepto la primera)
                if filename != 'index.md':
                    doc.add_page_break()
                
                process_markdown_file(doc, filepath)
        
        # Guardar documento
        output_path = os.path.join(docs_dir, 'Documentacion_Sistema_Deteccion_Tumores.docx')
        doc.save(output_path)
        
        print(f"\n✅ Documento generado exitosamente:")
        print(f"   📁 {output_path}")
        print(f"\n💡 Puedes abrirlo con Microsoft Word y guardarlo como PDF.")
        print(f"   En Word: Archivo > Guardar como > Tipo: PDF")
        
    except Exception as e:
        print(f"\n❌ Error al generar el documento: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
